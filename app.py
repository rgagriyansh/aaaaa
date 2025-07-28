from flask import Flask, render_template, request, send_file, jsonify, flash
import os
import tempfile
import uuid
from datetime import datetime
from werkzeug.utils import secure_filename
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'  # Change this in production

# Add health check route for debugging
@app.route('/health')
def health_check():
    return jsonify({
        "status": "healthy", 
        "timestamp": datetime.now().isoformat(),
        "claude_api_key_set": bool(os.getenv('CLAUDE_API_KEY'))
    })

# Configuration
UPLOAD_FOLDER = '/tmp'  # Use tmp directory on Vercel
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB max file size

# Note: No need to create upload folder - /tmp exists on Vercel

# Initialize the processor with better error handling
processor = None
try:
    # Temporarily disable to test basic Flask functionality
    # from main import IntegratedMathProcessor
    # processor = IntegratedMathProcessor()
    logger.info("Processor initialization skipped for testing")
except Exception as e:
    logger.error(f"Failed to initialize IntegratedMathProcessor: {e}")
    logger.error(f"CLAUDE_API_KEY present: {bool(os.getenv('CLAUDE_API_KEY'))}")
    # Don't crash the app, just log the error

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def generate_combined_document(extracted_texts):
    """
    Generate a single Word document containing all extracted texts with proper mathematical formatting
    
    Args:
        extracted_texts (list): List of dictionaries with 'filename' and 'text' keys
        
    Returns:
        str: Path to the generated document
    """
    try:
        from docx import Document
        from datetime import datetime
        import os
        
        # Generate filename first
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        final_filename = f"combined_math_document_{timestamp}.docx"
        
        # Create temporary documents for each section with proper math formatting
        temp_files = []
        all_expressions = []
        combined_text = ""
        
        try:
            # Process each text section individually to get proper math formatting
            for i, text_data in enumerate(extracted_texts, 1):
                filename = text_data['filename']
                text = text_data['text']
                
                if text.strip():
                    # Add section header
                    section_text = f"=== Image {i}: {filename} ===\n\n{text}"
                    combined_text += section_text + "\n\n"
                    
                    # Get mathematical expressions for this section
                    try:
                        expressions = processor.ask_openai_for_expressions(text)
                        if expressions:
                            all_expressions.extend(expressions)
                            logger.info(f"Found {len(expressions)} expressions in {filename}")
                    except Exception as e:
                        logger.warning(f"Error getting expressions for {filename}: {e}")
            
            # Now use the processor's unified document generation with all the combined text
            if all_expressions:
                logger.info(f"Generating combined document with {len(all_expressions)} total expressions")
                # Use the existing generate_unified_document method
                processor.generate_unified_document(combined_text, all_expressions, final_filename)
            else:
                # No mathematical expressions, create a simple document
                logger.info("No mathematical expressions found, creating simple document")
                doc = Document()
                doc.add_heading("Combined Math Document", level=1)
                doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph(f"Total images processed: {len(extracted_texts)}")
                doc.add_paragraph("")
                
                # Add content with section headers
                current_section = ""
                for line in combined_text.split('\n'):
                    if line.startswith('=== Image'):
                        # Save previous section
                        if current_section.strip():
                            doc.add_paragraph(current_section.strip())
                            current_section = ""
                        # Add new section header
                        header_text = line.replace('===', '').strip()
                        doc.add_heading(header_text, level=2)
                    else:
                        current_section += line + '\n'
                
                # Add final section
                if current_section.strip():
                    doc.add_paragraph(current_section.strip())
                
                doc.save(final_filename)
            
            return final_filename
            
        finally:
            # Clean up any temporary files
            for temp_file in temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                except Exception as e:
                    logger.warning(f"Error cleaning up temp file {temp_file}: {e}")
        
    except Exception as e:
        logger.error(f"Error generating combined document: {e}")
        raise

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_files():
    if processor is None:
        return jsonify({'error': 'Processor not available'}), 500
    
    if 'files' not in request.files:
        return jsonify({'error': 'No files selected'}), 400
    
    files = request.files.getlist('files')
    
    if not files or all(file.filename == '' for file in files):
        return jsonify({'error': 'No files selected'}), 400
    
    results = []
    temp_files = []
    all_extracted_texts = []
    successful_images = 0
    
    try:
        for file in files:
            if file and allowed_file(file.filename):
                # Create a unique filename
                filename = secure_filename(file.filename)
                unique_filename = f"{uuid.uuid4().hex}_{filename}"
                filepath = os.path.join(UPLOAD_FOLDER, unique_filename)
                
                # Save the uploaded file
                file.save(filepath)
                temp_files.append(filepath)
                
                try:
                    # Process the image
                    logger.info(f"Processing image: {filename}")
                    extracted_text, document_path, success = processor.process_image_to_document(filepath)
                    
                    if success:
                        successful_images += 1
                        all_extracted_texts.append({
                            'filename': filename,
                            'text': extracted_text
                        })
                        
                        results.append({
                            'filename': filename,
                            'status': 'success',
                            'extracted_text': extracted_text,
                            'message': 'Text extracted successfully'
                        })
                    else:
                        results.append({
                            'filename': filename,
                            'status': 'error',
                            'extracted_text': extracted_text if extracted_text else '',
                            'message': 'No mathematical expressions found or processing failed'
                        })
                        
                except Exception as e:
                    logger.error(f"Error processing {filename}: {e}")
                    results.append({
                        'filename': filename,
                        'status': 'error',
                        'extracted_text': '',
                        'message': f'Processing error: {str(e)}'
                    })
            else:
                results.append({
                    'filename': file.filename,
                    'status': 'error',
                    'extracted_text': '',
                    'message': 'Invalid file type'
                })
        
        # Generate combined document if we have successful extractions
        combined_document_path = None
        if successful_images > 0:
            try:
                combined_document_path = generate_combined_document(all_extracted_texts)
                logger.info(f"Combined document generated: {combined_document_path}")
            except Exception as e:
                logger.error(f"Error generating combined document: {e}")
                results.append({
                    'filename': 'Combined Document',
                    'status': 'error',
                    'extracted_text': '',
                    'message': f'Error creating combined document: {str(e)}'
                })
    
    finally:
        # Clean up uploaded files
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception as e:
                logger.error(f"Error cleaning up {temp_file}: {e}")
    
    return jsonify({
        'results': results,
        'combined_document_path': combined_document_path,
        'total_images': len(files),
        'successful_images': successful_images
    })

@app.route('/download/<path:filename>')
def download_file(filename):
    try:
        # Get the absolute path to the file
        file_path = os.path.abspath(filename)
        
        # Security check: ensure the file is in the current directory or subdirectories
        current_dir = os.path.abspath('.')
        if not file_path.startswith(current_dir):
            logger.error(f"Security violation: Attempted to access file outside current directory: {file_path}")
            return jsonify({'error': 'Access denied'}), 403
        
        # Check if file exists
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return jsonify({'error': 'File not found'}), 404
        
        # Get the original filename for download
        original_filename = os.path.basename(file_path)
        
        return send_file(file_path, as_attachment=True, download_name=original_filename)
    except Exception as e:
        logger.error(f"Error downloading file {filename}: {e}")
        return jsonify({'error': 'File not found'}), 404

# For Vercel serverless deployment
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000) 