import anthropic
import json
import re
import os
import sys
import glob
import base64
import logging
from datetime import datetime
from pathlib import Path
from PIL import Image
from docx import Document
from docx.oxml import OxmlElement

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('claude_responses.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Initialize Claude client
CLAUDE_API_KEY = os.getenv('CLAUDE_API_KEY')
if not CLAUDE_API_KEY:
    raise ValueError("CLAUDE_API_KEY environment variable is required")
client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

class IntegratedMathProcessor:
    def __init__(self, api_key=None):
        """
        Initialize the Integrated Math Processor with Claude API
        
        Args:
            api_key (str): Claude API key. If not provided, will use the default
        """
        if api_key:
            self.api_key = api_key
        else:
            self.api_key = CLAUDE_API_KEY
        
        if not self.api_key:
            raise ValueError("Claude API key is required.")
        
        self.client = anthropic.Anthropic(api_key=self.api_key)

    def encode_image(self, image_path):
        """
        Encode image to base64 for OpenAI API
        
        Args:
            image_path (str): Path to the image file
            
        Returns:
            tuple: (base64_string, media_type)
        """
        try:
            # Open and verify image
            with Image.open(image_path) as img:
                # Convert to RGB if necessary
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Save to bytes
                import io
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='JPEG')
                img_byte_arr = img_byte_arr.getvalue()
                
                # Encode to base64
                base64_string = base64.b64encode(img_byte_arr).decode('utf-8')
                return base64_string, "image/jpeg"
                
        except Exception as e:
            raise ValueError(f"Error processing image {image_path}: {str(e)}")

    def extract_text_from_image(self, image_path):
        """
        Extract and format text from image using OpenAI API
        
        Args:
            image_path (str): Path to the image file
            
        Returns:
            str: Extracted and formatted text
        """
        try:
            # Encode image
            base64_image, media_type = self.encode_image(image_path)
            
            # Create the prompt for OpenAI
            prompt = """Please analyze this image and identify ONLY the questions being asked. Do not include solutions, working steps, answers, or explanatory text.

Your task:
1. Look at the image and identify any questions or problems
2. Extract ONLY the question text (ignore any solutions, workings, or answers)
3. Format each question as a bullet point with •
4. Convert mathematical expressions to plain text format using ^ for exponents and / for fractions
5. Keep the natural flow of the question text
6. CRITICAL: ALWAYS put brackets around fractions when they appear in expressions

FRACTION FORMATTING RULES:
- Simple fractions like a/b should be written as (a/b)
- Complex fractions like (x+y)/z or x + (y/z) should keep their existing brackets
- When a fraction is part of a larger expression, bracket the fraction part
- Examples: n + 1/n should be written as n + (1/n), n^2 + 1/n^2 should be written as n^2 + (1/n^2)

Examples of the desired format:
• John has 2a^2 + 3b / 4c apples and Mary has x^2 + 5x - 3 toys
• The first equation is y = 3x^2 + 2x - 1 and the second is z = (a+b)/c
• If a, b, c are positive real numbers such that a + b + c = 12, find the minimum value of (a+1)(b+2)(c+3)
• The fraction (2x^2 + 3y) / (4z - 1) equals 5
• Solve for x in the equation (x^2 + 2x + 1) / (x + 1) = 3
• The expression (a^2 + b^2) / (2ab) represents the ratio
• If n + (1/n) = 3, find the values of n^2 + (1/n^2), n^4 + (1/n^4), and n^8 + (1/n^8)

IMPORTANT: Extract ONLY the questions. Do not include any solution steps, working, answers, or additional explanations. Focus solely on the question statements."""

            logger.info(f"[API CALL] Making Claude API call for image text extraction")
            logger.info(f"[PROMPT] Prompt sent to Claude: {prompt}")
            
            # Make API call to Claude
            response = self.client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=1024,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": media_type,
                                    "data": base64_image
                                }
                            },
                            {
                                "type": "text",
                                "text": prompt
                            }
                        ]
                    }
                ]
            )
            
            # Log the complete API response
            logger.info(f"[RESPONSE] Claude API Response Details:")
            logger.info(f"   Model: {response.model}")
            logger.info(f"   Usage: {response.usage}")
            logger.info(f"   Stop Reason: {response.stop_reason}")
            logger.info(f"   Raw Response Content: {response.content[0].text}")
            
            extracted_text = response.content[0].text
            
            logger.info(f"[SUCCESS] Successfully extracted text from image")
            logger.info(f"[TEXT] Extracted text length: {len(extracted_text)} characters")
            logger.info(f"[TEXT] Extracted text preview: {extracted_text[:200]}{'...' if len(extracted_text) > 200 else ''}")
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"[ERROR] Error extracting text from image: {str(e)}")
            raise RuntimeError(f"Error extracting text from image: {str(e)}")

    def format_superscripts(self, expr):
        """Convert ^digit to Unicode superscripts (basic handling)"""
        superscripts = {'0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
                        '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹'}
        return re.sub(r'\^(\d)', lambda m: superscripts.get(m.group(1), m.group(1)), expr)

    def find_outermost_division(self, expr):
        """
        Find the position of the division operator that's not inside parentheses
        Returns -1 if no valid division found
        """
        paren_count = 0
        for i, char in enumerate(expr):
            if char == '(':
                paren_count += 1
            elif char == ')':
                paren_count -= 1
            elif char == '/' and paren_count == 0:
                return i
        return -1

    def remove_outer_parens(self, expr):
        """
        Remove outer parentheses if they exist and are balanced
        """
        expr = expr.strip()
        if expr.startswith('(') and expr.endswith(')'):
            # Check if parentheses are balanced
            paren_count = 0
            for i, char in enumerate(expr):
                if char == '(':
                    paren_count += 1
                elif char == ')':
                    paren_count -= 1
                    if paren_count == 0 and i == len(expr) - 1:
                        # Balanced parentheses, remove them
                        return expr[1:-1].strip()
                    elif paren_count < 0:
                        # Unbalanced, don't remove
                        break
        return expr

    def create_math_equation(self, paragraph, expr):
        """Creates proper mathematical formatting - handles mixed expressions with fractions"""
        
        # Check if this is a pure fraction (like (1/n) or a/b)
        if self.is_pure_fraction(expr):
            # Handle as a fraction
            self.create_fraction_element(paragraph, expr)
        else:
            # Handle as a mixed expression that may contain fractions
            self.create_mixed_expression_element(paragraph, expr)
    
    def is_pure_fraction(self, expr):
        """Check if expression is a pure fraction (like (1/n) or a/b)"""
        expr = expr.strip()
        
        # Handle modulus functions |...|
        if expr.startswith('|') and expr.endswith('|'):
            inner_expr = expr[1:-1].strip()
            # Check if the inner expression is a pure fraction
            return self.is_pure_fraction(inner_expr)
        
        # Handle other function wrappers
        func_match = re.match(r'([a-zA-Z]+)\((.*)\)', expr)
        if func_match:
            inner_expr = func_match.group(2)
            # Check if the inner expression is a pure fraction
            return self.is_pure_fraction(inner_expr)
        
        # If it's wrapped in parentheses and contains exactly one /, it's likely a fraction
        if expr.startswith('(') and expr.endswith(')') and expr.count('/') == 1:
            return True
        
        # If it contains exactly one / and no other operators, it's a simple fraction
        if expr.count('/') == 1:
            # Check if there are other operators
            operators = ['+', '-', '*', '=']
            for op in operators:
                if op in expr:
                    return False
            return True
        
        return False
    
    def should_be_single_math_expression(self, expr):
        """Check if expression should be treated as a single mathematical expression"""
        expr = expr.strip()
        
        # If it contains mathematical operators and fractions, it's likely a single expression
        math_operators = ['+', '-', '*', '/', '=']
        has_operators = any(op in expr for op in math_operators)
        has_fractions = '(' in expr and ')' in expr and '/' in expr
        
        # Examples that should be single expressions:
        # - n + (1/n) = 3
        # - n^2 + (1/n^2)
        # - x + (1/x) = 5
        # - a + b/(c+d)
        
        if has_operators and has_fractions:
            # Check if it's a coherent mathematical expression
            # Look for patterns like: variable + (fraction) or variable + (fraction) = value
            if re.search(r'[a-zA-Z]\s*\+\s*\([^)]*\/[^)]*\)', expr):
                return True
            if re.search(r'[a-zA-Z]\s*\+\s*\([^)]*\/[^)]*\)\s*=', expr):
                return True
            if re.search(r'[a-zA-Z]\^[0-9]\s*\+\s*\([^)]*\/[^)]*\)', expr):
                return True
        
        return False
    
    def create_single_math_with_inline_fractions(self, paragraph, expr):
        """Create a single math element with inline fractions"""
        # Format superscripts first
        formatted_expr = self.format_superscripts(expr.strip())
        
        # Create the main math container
        oMath = OxmlElement('m:oMath')
        
        # Split the expression into parts: text and fractions
        parts = self.extract_fractions_from_expression(formatted_expr)
        
        if not parts:
            # No fractions found, create simple math element
            math_r = OxmlElement('m:r')
            math_t = OxmlElement('m:t')
            math_t.text = formatted_expr
            math_r.append(math_t)
            oMath.append(math_r)
        else:
            # Process each part
            for part_type, part_text in parts:
                if part_type == 'text':
                    # Add text as a math run
                    if part_text.strip():
                        math_r = OxmlElement('m:r')
                        math_t = OxmlElement('m:t')
                        math_t.text = part_text.strip()
                        math_r.append(math_t)
                        oMath.append(math_r)
                elif part_type == 'fraction':
                    # Create fraction element
                    self.create_inline_fraction(oMath, part_text)
        
        # Create the math paragraph
        oMathPara = OxmlElement('m:oMathPara')
        oMathPara.append(oMath)
        paragraph._element.append(oMathPara)
    
    def create_inline_fraction(self, oMath, expr):
        """Create an inline fraction element"""
        # Handle modulus functions |...|
        if expr.startswith('|') and expr.endswith('|'):
            # Create modulus function with fraction inside
            self.create_modulus_with_fraction(oMath, expr)
            return
        
        # Handle other function wrappers like sqrt(...), log(...), etc.
        func_match = re.match(r'([a-zA-Z]+)\((.*)\)', expr)
        if func_match:
            func_name = func_match.group(1)
            inner_expr = func_match.group(2)
            # Create function with fraction inside
            self.create_function_with_fraction(oMath, func_name, inner_expr)
            return
        
        # Remove outer parentheses if they exist
        if expr.startswith('(') and expr.endswith(')'):
            expr = expr[1:-1].strip()
        
        # Split by division
        parts = expr.split('/')
        if len(parts) != 2:
            # Fallback to simple text if parsing fails
            math_r = OxmlElement('m:r')
            math_t = OxmlElement('m:t')
            math_t.text = expr
            math_r.append(math_t)
            oMath.append(math_r)
            return
        
        numerator = parts[0].strip()
        denominator = parts[1].strip()
        
        # Create the fraction
        frac = OxmlElement('m:f')
        
        num = OxmlElement('m:num')
        num_r = OxmlElement('m:r')
        num_t = OxmlElement('m:t')
        num_t.text = numerator
        num_r.append(num_t)
        num.append(num_r)
        
        den = OxmlElement('m:den')
        den_r = OxmlElement('m:r')
        den_t = OxmlElement('m:t')
        den_t.text = denominator
        den_r.append(den_t)
        den.append(den_r)
        
        frac.append(num)
        frac.append(den)
        oMath.append(frac)
    
    def create_modulus_with_fraction(self, oMath, expr):
        """Create a modulus function with a fraction inside"""
        # Remove the modulus bars
        inner_expr = expr[1:-1].strip()
        
        # Check if the inner expression contains a division operator (fraction)
        if '/' in inner_expr:
            # Create the modulus function with the entire inner expression
            self.create_modulus_function(oMath, inner_expr)
        else:
            # If not a fraction, create simple modulus
            self.create_modulus_function(oMath, inner_expr)
    
    def create_modulus_function(self, oMath, inner_expr):
        """Create a modulus function element"""
        # Check if the argument is a fraction
        if '/' in inner_expr:
            # Create the modulus function that wraps the entire fraction
            # We'll create it as |(fraction)| format using direct bars
            
            # Add opening modulus bar
            open_bar = OxmlElement('m:r')
            open_bar_t = OxmlElement('m:t')
            open_bar_t.text = '|'
            open_bar.append(open_bar_t)
            oMath.append(open_bar)
            
            # Create the fraction element
            frac = OxmlElement('m:f')
            
            # Split by division
            parts = inner_expr.split('/')
            if len(parts) == 2:
                numerator = parts[0].strip()
                denominator = parts[1].strip()
                
                # Remove outer parentheses if they exist
                if numerator.startswith('(') and numerator.endswith(')'):
                    numerator = numerator[1:-1].strip()
                if denominator.startswith('(') and denominator.endswith(')'):
                    denominator = denominator[1:-1].strip()
                
                # Create numerator
                num = OxmlElement('m:num')
                num_r = OxmlElement('m:r')
                num_t = OxmlElement('m:t')
                num_t.text = numerator
                num_r.append(num_t)
                num.append(num_r)
                
                # Create denominator
                den = OxmlElement('m:den')
                den_r = OxmlElement('m:r')
                den_t = OxmlElement('m:t')
                den_t.text = denominator
                den_r.append(den_t)
                den.append(den_r)
                
                frac.append(num)
                frac.append(den)
                oMath.append(frac)
            
            # Add closing modulus bar
            close_bar = OxmlElement('m:r')
            close_bar_t = OxmlElement('m:t')
            close_bar_t.text = '|'
            close_bar.append(close_bar_t)
            oMath.append(close_bar)
        else:
            # Create simple modulus function for non-fraction expressions
            # Add opening modulus bar
            open_bar = OxmlElement('m:r')
            open_bar_t = OxmlElement('m:t')
            open_bar_t.text = '|'
            open_bar.append(open_bar_t)
            oMath.append(open_bar)
            
            # Add the inner expression
            arg_r = OxmlElement('m:r')
            arg_t = OxmlElement('m:t')
            arg_t.text = inner_expr
            arg_r.append(arg_t)
            oMath.append(arg_r)
            
            # Add closing modulus bar
            close_bar = OxmlElement('m:r')
            close_bar_t = OxmlElement('m:t')
            close_bar_t.text = '|'
            close_bar.append(close_bar_t)
            oMath.append(close_bar)
    
    def create_function_with_fraction(self, oMath, func_name, inner_expr):
        """Create a function with a fraction inside"""
        # Create the function element
        func = OxmlElement('m:func')
        
        # Function name
        func_name_elem = OxmlElement('m:fName')
        func_name_r = OxmlElement('m:r')
        func_name_t = OxmlElement('m:t')
        func_name_t.text = func_name
        func_name_r.append(func_name_t)
        func_name_elem.append(func_name_r)
        
        # Function argument
        func_arg = OxmlElement('m:e')
        
        # Check if the argument is a fraction
        if '(' in inner_expr and ')' in inner_expr and '/' in inner_expr:
            # Create fraction element
            self.create_inline_fraction(func_arg, inner_expr)
        else:
            # Create simple math element
            arg_r = OxmlElement('m:r')
            arg_t = OxmlElement('m:t')
            arg_t.text = inner_expr
            arg_r.append(arg_t)
            func_arg.append(arg_r)
        
        func.append(func_name_elem)
        func.append(func_arg)
        oMath.append(func)
    
    def extract_fractions_from_expression(self, expr):
        """Extract fractions from a mixed expression and return parts to format"""
        parts = []
        current_pos = 0
        
        # Find all fractions in the expression
        while current_pos < len(expr):
            # Look for fractions with proper handling of nested parentheses
            fraction_info = self.find_fraction_with_nested_parens(expr[current_pos:])
            
            if fraction_info:
                start = current_pos + fraction_info['start']
                end = current_pos + fraction_info['end']
                fraction_text = fraction_info['text']
                
                # Check if this fraction is wrapped in modulus or other functions
                wrapper_info = self.find_function_wrapper(expr, start, end)
                
                if wrapper_info:
                    wrapper_start = wrapper_info['start']
                    wrapper_end = wrapper_info['end']
                    wrapped_text = wrapper_info['text']
                    
                    # Add text before the wrapped fraction
                    if wrapper_start > current_pos:
                        text_before = expr[current_pos:wrapper_start]
                        if text_before.strip():
                            parts.append(('text', text_before))
                    
                    # Add the wrapped fraction
                    parts.append(('fraction', wrapped_text))
                    current_pos = wrapper_end
                else:
                    # Add text before the fraction
                    if start > current_pos:
                        text_before = expr[current_pos:start]
                        if text_before.strip():
                            parts.append(('text', text_before))
                    
                    # Add the fraction
                    parts.append(('fraction', fraction_text))
                    current_pos = end
            else:
                # No more fractions found, add remaining text
                remaining_text = expr[current_pos:]
                if remaining_text.strip():
                    parts.append(('text', remaining_text))
                break
        
        return parts
    
    def find_fraction_with_nested_parens(self, expr):
        """Find a fraction with proper handling of nested parentheses"""
        # Look for division operator first
        div_pos = expr.find('/')
        if div_pos == -1:
            return None
        
        # Find the start of the fraction (look for opening parenthesis before the division)
        start_pos = -1
        for i in range(div_pos - 1, -1, -1):
            if expr[i] == ')':
                # Skip to the matching opening parenthesis
                paren_count = 0
                for j in range(i, -1, -1):
                    if expr[j] == ')':
                        paren_count += 1
                    elif expr[j] == '(':
                        paren_count -= 1
                        if paren_count == 0:
                            start_pos = j
                            break
                if start_pos != -1:
                    break
            elif expr[i] == '(':
                start_pos = i
                break
            elif expr[i].isspace():
                continue
            else:
                # Found a character that's not a closing parenthesis, space, or opening parenthesis
                # This means we're looking at the numerator part
                # Find the start of this term
                for j in range(i, -1, -1):
                    if expr[j] in '+-*/()':
                        start_pos = j + 1
                        break
                if start_pos == -1:
                    start_pos = 0
                break
        
        if start_pos == -1:
            start_pos = 0
        
        # Find the end of the fraction (look for closing parenthesis after the division)
        end_pos = -1
        for i in range(div_pos + 1, len(expr)):
            if expr[i] == '(':
                # Skip to the matching closing parenthesis
                paren_count = 0
                for j in range(i, len(expr)):
                    if expr[j] == '(':
                        paren_count += 1
                    elif expr[j] == ')':
                        paren_count -= 1
                        if paren_count == 0:
                            end_pos = j + 1
                            break
                if end_pos != -1:
                    break
            elif expr[i] == ')':
                end_pos = i + 1
                break
            elif expr[i].isspace():
                continue
            else:
                # Found a character that's not an opening parenthesis, space, or closing parenthesis
                # This means we're looking at the denominator part
                # Find the end of this term
                for j in range(i, len(expr)):
                    if expr[j] in '+-*/()':
                        end_pos = j
                        break
                if end_pos == -1:
                    end_pos = len(expr)
                break
        
        if end_pos == -1:
            end_pos = len(expr)
        
        # Extract the fraction
        fraction_text = expr[start_pos:end_pos]
        
        # Verify it's actually a fraction (contains division operator)
        if '/' in fraction_text:
            return {
                'start': start_pos,
                'end': end_pos,
                'text': fraction_text
            }
        
        return None
    
    def find_function_wrapper(self, expr, start, end):
        """Find if a fraction is wrapped in a function like |...| or sqrt(...)"""
        # Check for modulus function |...| that wraps the entire fraction
        # Look for opening | before the fraction
        open_pos = -1
        for i in range(start - 1, -1, -1):
            if expr[i] == '|':
                open_pos = i
                break
            elif expr[i].isspace():
                continue
            else:
                # Found a character that's not | or space, stop looking
                break
        
        if open_pos != -1:
            # Look for closing | after the fraction
            close_pos = -1
            for i in range(end, len(expr)):
                if expr[i] == '|':
                    close_pos = i
                    break
                elif expr[i].isspace():
                    continue
                else:
                    # Found a character that's not | or space, stop looking
                    break
            
            if close_pos != -1:
                return {
                    'start': open_pos,
                    'end': close_pos + 1,
                    'text': expr[open_pos:close_pos+1]
                }
        
        # Check for other function wrappers (sqrt, log, etc.)
        # Look for function names followed by opening parenthesis
        func_pattern = r'([a-zA-Z]+)\('
        func_match = re.search(func_pattern, expr[:start])
        if func_match:
            func_name = func_match.group(1)
            # Find the closing parenthesis for this function
            paren_count = 0
            for i in range(func_match.end()-1, len(expr)):
                if expr[i] == '(':
                    paren_count += 1
                elif expr[i] == ')':
                    paren_count -= 1
                    if paren_count == 0:
                        return {
                            'start': func_match.start(),
                            'end': i + 1,
                            'text': expr[func_match.start():i+1]
                        }
        
        return None
    
    def create_fraction_element(self, paragraph, expr):
        """Create a fraction element for pure fractions"""
        # Remove outer parentheses if they exist
        if expr.startswith('(') and expr.endswith(')'):
            expr = expr[1:-1].strip()
        
        # Split by division
        parts = expr.split('/')
        if len(parts) != 2:
            # Fallback to simple math if parsing fails
            self.create_simple_math_element(paragraph, expr)
            return
        
        numerator = self.format_superscripts(parts[0].strip())
        denominator = self.format_superscripts(parts[1].strip())
        
        # Create the fraction
        frac = OxmlElement('m:f')
        
        num = OxmlElement('m:num')
        num_r = OxmlElement('m:r')
        num_t = OxmlElement('m:t')
        num_t.text = numerator
        num_r.append(num_t)
        num.append(num_r)
        
        den = OxmlElement('m:den')
        den_r = OxmlElement('m:r')
        den_t = OxmlElement('m:t')
        den_t.text = denominator
        den_r.append(den_t)
        den.append(den_r)
        
        frac.append(num)
        frac.append(den)
        
        oMath = OxmlElement('m:oMath')
        oMath.append(frac)
        oMathPara = OxmlElement('m:oMathPara')
        oMathPara.append(oMath)
        paragraph._element.append(oMathPara)
    
    def create_mixed_expression_element(self, paragraph, expr):
        """Create a mixed expression element that may contain fractions"""
        # Check if this should be treated as a single mathematical expression
        if self.should_be_single_math_expression(expr):
            # Create as a single math element with inline fractions
            self.create_single_math_with_inline_fractions(paragraph, expr)
        else:
            # Extract fractions from the expression
            parts = self.extract_fractions_from_expression(expr)
            
            if not parts:
                # No fractions found, create simple math element
                self.create_simple_math_element(paragraph, expr)
                return
            
            # Create a mixed expression with proper fraction formatting
            for part_type, part_text in parts:
                if part_type == 'text':
                    # Add regular text with superscript formatting
                    if part_text.strip():
                        # Format superscripts in the text part
                        formatted_text = self.format_superscripts(part_text.strip())
                        run = paragraph.add_run(formatted_text)
                elif part_type == 'fraction':
                    # Create fraction element
                    self.create_fraction_element(paragraph, part_text)
    
    def create_simple_math_element(self, paragraph, expr):
        """Create a simple math expression element"""
        formatted_expr = self.format_superscripts(expr.strip())
        
        # Create simple math run
        math_r = OxmlElement('m:r')
        math_t = OxmlElement('m:t')
        math_t.text = formatted_expr
        math_r.append(math_t)
        
        oMath = OxmlElement('m:oMath')
        oMath.append(math_r)
        oMathPara = OxmlElement('m:oMathPara')
        oMathPara.append(oMath)
        paragraph._element.append(oMathPara)

    def ask_openai_for_expressions(self, text):
        """Use OpenAI to identify and extract mathematical expressions from text"""
        
        prompt = f"""
        Analyze the following text and identify ALL mathematical expressions in it. For each mathematical expression found, extract:
        1. The text before the expression (context/statement)
        2. The mathematical expression itself
        3. The text after the expression (if any)
        
        Text to analyze: "{text}"
        
        Please respond with a JSON array where each object has this format:
        {{
            "statement": "text before the expression",
            "expression": "the mathematical expression using ^ for superscripts",
            "suffix": "text after the expression",
            "original_context": "the full sentence containing this expression"
        }}
        
        Mathematical expressions include:
        - Variables with coefficients (2x, 3a)
        - Superscripts (x^2, a^3)
        - Fractions (a/b, (x+y)/z, (2x^2 + 3y)/(4z - 1))
        - Operations (+, -, *, /)
        - Parentheses for grouping and complex expressions

       
        
        CRITICAL FRACTION FORMATTING RULES:
        1. ALWAYS put brackets around fractions when they appear in expressions
        2. Simple fractions like a/b should be written as (a/b)
        3. Complex fractions like (x+y)/z should keep their existing brackets
        4. When a fraction is part of a larger expression, bracket the fraction part
        5. Examples of correct formatting:
           - n + 1/n should be written as n + (1/n)
           - n^2 + 1/n^2 should be written as n^2 + (1/n^2)
           - x + (1/x) = 3 should be written as x + (1/x) = 3
           - (a^2 + b^2) / (2ab) should keep its brackets
        
        Examples of correct expressions:
        - "2a^2 + 3b" 
        - "x^2 + 5x - 3"
        - "(a^2 + b^2) / 2c"
        - "(2x^2 + 3y) / (4z - 1)"
        - "(x^2 + 2x + 1) / (x + 1)"
        - "(a^2 + b^2) / (2ab)"
        - "(a+1)(b+2)(c+3)"
        - "x + (1/x) = 3"
        - "n + (1/n) = 3"
        - "n^2 + (1/n^2)"
        - "n^4 + (1/n^4)"
        - "n^8 + (1/n^8)"
        
        If no mathematical expressions are found, return an empty array [].
        Return only the JSON, no other text.
        """
        
        logger.info(f"[API CALL] Making Claude API call for mathematical expression identification")
        logger.info(f"[TEXT] Text to analyze: {text[:200]}{'...' if len(text) > 200 else ''}")
        logger.info(f"[PROMPT] Prompt sent to Claude: {prompt}")
        
        try:
            message = self.client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=1000,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            # Log the complete API response
            logger.info(f"[RESPONSE] Claude API Response Details:")
            logger.info(f"   Model: {message.model}")
            logger.info(f"   Usage: {message.usage}")
            logger.info(f"   Stop Reason: {message.stop_reason}")
            logger.info(f"   Raw Response Content: {message.content[0].text}")
            
            # Extract JSON from response
            response_text = message.content[0].text.strip()
            
            # Try to parse JSON
            try:
                expressions = json.loads(response_text)
                logger.info(f"[SUCCESS] Successfully parsed JSON response")
                logger.info(f"[COUNT] Found {len(expressions)} mathematical expressions")
                for i, expr in enumerate(expressions, 1):
                    logger.info(f"   Expression {i}: {expr}")
                return expressions
            except json.JSONDecodeError as e:
                logger.warning(f"[WARNING] Direct JSON parsing failed: {e}")
                logger.warning(f"[WARNING] Raw response that failed to parse: {response_text}")
                
                # If direct parsing fails, try to extract JSON from the response
                json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
                if json_match:
                    try:
                        expressions = json.loads(json_match.group())
                        logger.info(f"[SUCCESS] Successfully extracted JSON from response using regex")
                        logger.info(f"[COUNT] Found {len(expressions)} mathematical expressions")
                        for i, expr in enumerate(expressions, 1):
                            logger.info(f"   Expression {i}: {expr}")
                        return expressions
                    except json.JSONDecodeError as e2:
                        logger.error(f"[ERROR] Regex JSON extraction also failed: {e2}")
                        logger.error(f"[ERROR] Extracted text that failed: {json_match.group()}")
                        print(f"[WARNING] Claude response was not valid JSON: {response_text}")
                        return []
                else:
                    logger.error(f"[ERROR] No JSON array found in response")
                    print(f"[WARNING] Claude response was not valid JSON: {response_text}")
                    return []
                    
        except Exception as e:
            logger.error(f"[ERROR] Error calling Claude API: {str(e)}")
            print(f"[ERROR] Error calling Claude API: {str(e)}")
            return []

    def generate_unified_document(self, original_text, expressions, output_path):
        """Generate a single Word document with all expressions formatted in their original context"""
        doc = Document()
        doc.add_heading("Math Expression Document", level=1)

        # Find all mathematical expressions in the original text and their positions
        expression_positions = []
        
        for expr_data in expressions:
            expr = expr_data['expression']
            # Find all occurrences of this expression in the original text
            start = 0
            while True:
                pos = original_text.find(expr, start)
                if pos == -1:
                    break
                
                # Check if this position hasn't been used yet
                if not any(ep['start'] <= pos < ep['end'] for ep in expression_positions):
                    expression_positions.append({
                        'start': pos,
                        'end': pos + len(expr),
                        'expression': expr,
                        'data': expr_data
                    })
                    break
                start = pos + 1
        
        # Sort by position (ascending to process from beginning to end)
        expression_positions.sort(key=lambda x: x['start'])
        
        # Create the document paragraph
        p = doc.add_paragraph()
        
        # Build the document by processing the original text
        last_pos = 0
        
        for expr_pos in expression_positions:
            # Add text before this expression
            text_before = original_text[last_pos:expr_pos['start']]
            if text_before:
                p.add_run(text_before)
            
            # Add the formatted mathematical expression
            self.create_math_equation(p, expr_pos['expression'])
            
            # Update position
            last_pos = expr_pos['end']
        
        # Add any remaining text after the last expression
        remaining_text = original_text[last_pos:]
        if remaining_text:
            p.add_run(remaining_text)

        doc.save(output_path)
        return output_path

    def process_image_to_document(self, image_path):
        """
        Complete flow: Extract text from image (for web app - no individual document generation)
        
        Args:
            image_path (str): Path to the image file
            
        Returns:
            tuple: (extracted_text, document_path, success)
        """
        try:
            logger.info(f"[START] Starting image processing flow for: {image_path}")
            
            # Step 1: Extract text from image
            print(f"[IMAGE] Processing image: {os.path.basename(image_path)}")
            print("[OCR] Extracting questions from image...")
            logger.info(f"[STEP1] Step 1: Extracting text from image using ChatGPT")
            
            extracted_text = self.extract_text_from_image(image_path)
            
            if not extracted_text.strip():
                logger.warning(f"[WARNING] No text extracted from image")
                return extracted_text, None, False
            
            print("[SUCCESS] Questions extracted successfully!")
            print(f"[TEXT] Extracted text: {extracted_text[:100]}{'...' if len(extracted_text) > 100 else ''}")
            logger.info(f"[SUCCESS] Step 1 completed: Text extraction successful")
            
            # For web app, we only extract text and return success
            # The combined document will be generated separately
            logger.info(f"[COMPLETE] Text extraction completed successfully")
            
            return extracted_text, None, True
            
        except Exception as e:
            logger.error(f"[ERROR] Error in processing flow: {str(e)}")
            print(f"[ERROR] Error in processing: {str(e)}")
            return "", None, False

def display_menu():
    """Display the main menu"""
    print("\n" + "="*60)
    print("        INTEGRATED MATH PROCESSOR")
    print("     Image OCR → Formatted Math Documents (Claude)")
    print("="*60)
    print("1. 🖼️  Process single image → Generate document")
    print("2. 📁 Process multiple images → Generate documents")
    print("3. 📝 Enter text directly → Generate document")
    print("4. 🔍 Extract questions only (no document)")
    print("5. ❌ Exit")
    print("-"*60)

def get_supported_images(directory):
    """Get all supported image files from directory"""
    supported_extensions = ['*.jpg', '*.jpeg', '*.png', '*.bmp', '*.tiff', '*.webp']
    image_files = []
    
    for ext in supported_extensions:
        image_files.extend(glob.glob(os.path.join(directory, ext)))
        image_files.extend(glob.glob(os.path.join(directory, ext.upper())))
    
    return sorted(image_files)

def process_single_image_to_document(processor):
    """Process a single image and generate formatted document"""
    image_path = input("\n📎 Enter the path to your image file: ").strip().strip('"')
    
    if not os.path.exists(image_path):
        print(f"❌ Error: File '{image_path}' not found!")
        return
    
    try:
        extracted_text, document_path, success = processor.process_image_to_document(image_path)
        
        if success:
            print("\n" + "="*60)
            print("🎉 SUCCESS! Complete processing finished")
            print("="*60)
            print(f"📝 Extracted Questions:")
            print("-" * 40)
            print(extracted_text)
            print("-" * 40)
            print(f"📄 Generated Document: {document_path}")
            print("="*60)
        else:
            print("\n❌ Processing failed or no mathematical expressions found")
            if extracted_text:
                print("📝 Extracted text:")
                print(extracted_text)
        
    except Exception as e:
        print(f"❌ Error processing image: {str(e)}")

def process_multiple_images_to_documents(processor):
    """Process multiple images and generate formatted documents"""
    directory = input("\n📁 Enter the directory path containing images: ").strip().strip('"')
    
    if not os.path.exists(directory):
        print(f"❌ Error: Directory '{directory}' not found!")
        return
    
    # Find all image files
    image_files = get_supported_images(directory)
    
    if not image_files:
        print(f"❌ No supported image files found in '{directory}'")
        print("Supported formats: JPG, JPEG, PNG, BMP, TIFF, WEBP")
        return
    
    print(f"\n📊 Found {len(image_files)} image file(s):")
    for i, img in enumerate(image_files, 1):
        print(f"  {i}. {os.path.basename(img)}")
    
    # Process each image
    results = []
    successful_docs = 0
    
    print(f"\n🚀 Starting batch processing...")
    
    for i, image_path in enumerate(image_files, 1):
        print(f"\n{'='*60}")
        print(f"Processing {i}/{len(image_files)}: {os.path.basename(image_path)}")
        print('='*60)
        
        try:
            extracted_text, document_path, success = processor.process_image_to_document(image_path)
            
            if success:
                successful_docs += 1
                results.append((image_path, "✅ Success", document_path, extracted_text))
            else:
                results.append((image_path, "❌ Failed", None, extracted_text))
                
        except Exception as e:
            print(f"❌ Error processing {os.path.basename(image_path)}: {str(e)}")
            results.append((image_path, f"❌ Error: {str(e)}", None, ""))
    
    # Display final summary
    print("\n" + "="*80)
    print("📊 BATCH PROCESSING SUMMARY")
    print("="*80)
    print(f"Total files processed: {len(image_files)}")
    print(f"Successful documents: {successful_docs}")
    print(f"Failed/No math expressions: {len(image_files) - successful_docs}")
    print("\nDetails:")
    print("-" * 80)
    
    for image_path, status, doc_path, text in results:
        print(f"\n📄 {os.path.basename(image_path)}")
        print(f"   Status: {status}")
        if doc_path:
            print(f"   Document: {doc_path}")
        if text and len(text) > 0:
            preview = text[:100] + "..." if len(text) > 100 else text
            print(f"   Preview: {preview}")

def process_text_directly(processor):
    """Process text directly without OCR"""
    print("\n📝 Direct Text Processing")
    print("=" * 50)
    print("Enter text containing mathematical expressions.")
    print("The AI will automatically detect and format all expressions!")
    print()
    print("Examples:")
    print("• John has 2a^2 + 3b / 4c apples and Mary has x^2 + 5x - 3 toys")
    print("• The first equation is y = 3x^2 + 2x - 1 and the second is z = (a+b)/c")
    print("• The fraction (2x^2 + 3y) / (4z - 1) equals 5")
    print("• Solve for x in the equation (x^2 + 2x + 1) / (x + 1) = 3")
    print("• The expression (a^2 + b^2) / (2ab) represents the ratio")
    print()
    
    input_text = input("Your text: ").strip()
    
    if not input_text:
        print("❌ No input provided!")
        return
    
    try:
        print(f"\n🔍 Analyzing text: {input_text}")
        print("🤖 Identifying mathematical expressions...")
        
        expressions = processor.ask_openai_for_expressions(input_text)
        
        if not expressions:
            print("❌ No mathematical expressions detected!")
            return
        
        print(f"\n✅ Found {len(expressions)} mathematical expression(s)!")
        
        for i, expr_data in enumerate(expressions, 1):
            print(f"   {i}. 📝 '{expr_data['statement']}' 🔢 [{expr_data['expression']}] 📝 '{expr_data['suffix']}'")
        
        # Generate document
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"math_doc_text_{timestamp}.docx"
        
        print(f"\n📄 Creating document: {filename}")
        
        output_path = processor.generate_unified_document(input_text, expressions, filename)
        
        print(f"\n🎉 Success! Document created: {os.path.abspath(output_path)}")
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")

def extract_questions_only(processor):
    """Extract questions from image without generating document"""
    image_path = input("\n📎 Enter the path to your image file: ").strip().strip('"')
    
    if not os.path.exists(image_path):
        print(f"❌ Error: File '{image_path}' not found!")
        return
    
    try:
        print(f"🖼️  Processing image: {os.path.basename(image_path)}")
        print("📖 Extracting questions from image...")
        
        extracted_text = processor.extract_text_from_image(image_path)
        
        print("\n" + "="*60)
        print("📝 EXTRACTED QUESTIONS:")
        print("="*60)
        print(extracted_text)
        print("="*60)
        
        # Ask if user wants to save
        save_choice = input("\n💾 Save extracted text to file? (y/n): ").lower().strip()
        if save_choice == 'y':
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            image_name = Path(image_path).stem
            filename = f"extracted_questions_{image_name}_{timestamp}.txt"
            
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(f"Questions extracted from: {image_path}\n")
                f.write("=" * 50 + "\n\n")
                f.write(extracted_text)
            
            print(f"✅ Saved to: {os.path.abspath(filename)}")
        
    except Exception as e:
        print(f"❌ Error extracting text: {str(e)}")

def main():
    """Main interactive function"""
    try:
        logger.info(f"[START] Starting Integrated Math Processor application")
        
        # Initialize processor
        processor = IntegratedMathProcessor()
        
        print("[START] Integrated Math Processor (Claude) initialized successfully!")
        logger.info(f"[SUCCESS] Integrated Math Processor initialized successfully")
        
        while True:
            display_menu()
            choice = input("Enter your choice (1-5): ").strip()
            
            if choice == '1':
                process_single_image_to_document(processor)
            
            elif choice == '2':
                process_multiple_images_to_documents(processor)
            
            elif choice == '3':
                process_text_directly(processor)
            
            elif choice == '4':
                extract_questions_only(processor)
            
            elif choice == '5':
                print("\n👋 Thank you for using Integrated Math Processor (Claude)!")
                break
            
            else:
                print("\n❌ Invalid choice! Please enter 1, 2, 3, 4, or 5.")
            
            # Ask if user wants to continue
            if choice in ['1', '2', '3', '4']:
                continue_choice = input("\n🔄 Process more content? (y/n): ").lower().strip()
                if continue_choice != 'y':
                    print("\n👋 Thank you for using Integrated Math Processor (Claude)!")
                    break
    
    except KeyboardInterrupt:
        print("\n\n⏹️  Operation cancelled by user.")
    except Exception as e:
        print(f"\n❌ Unexpected error: {str(e)}")

if __name__ == "__main__":
    main() 